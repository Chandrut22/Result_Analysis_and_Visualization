

# Create your views here.
from django.contrib.auth.forms import UserCreationForm
from django.shortcuts import render,redirect
from django.http import HttpResponse
from .forms import CreateUserForm
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.decorators import login_required
import tabula
import pandas as pd
import numpy as np
import io
import os
from django import forms
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
import seaborn as sns
import base64
import docx
from docx.shared import Inches
from docx import Document
from io import BytesIO
import PyPDF2
from PyPDF2 import PdfReader

#Overall Performance Function
def overall_per(request):
    data = pd.read_csv('E:\sra\Student_Result.csv')
    df = data.copy()
    df.drop('Unnamed: 0',axis=1,inplace=True)
    df["Arrears"] = df.isin([0]).sum(axis=1)
    uni_arr = list(df['Arrears'].unique())
    label = []
    plot = []
    arrears_data = pd.DataFrame()
    for i in range(0, len(uni_arr)):
            value = int((df['Arrears'] == uni_arr[i]).sum())
            plot.append(value)
            label.append(str(value)+" Students Have "+str(uni_arr[i])+" arrears.")
            
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.pie(plot, labels=label, startangle=90, autopct='%1.1f%%')
    ax.axis("equal")
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    plt.close(fig)
    plot_data = buffer.getvalue()
    encoded_image = base64.b64encode(plot_data).decode('utf-8')
        
    doc = docx.Document()
    doc.add_heading("Overall Performance", 0)
    doc.add_heading("Performance Chart:",level=2)
    doc.add_picture(io.BytesIO(base64.b64decode(encoded_image)), width=Inches(6))
    doc.add_heading("Performance Summary:",level=2)
        
    dic = {}
    uni_arr.sort(reverse=True)
    for i in range(0, len(uni_arr)):
            key = "Student Have "+str(uni_arr[i])+" arrears:"
            doc.add_heading(key,level=3)
            temp=df[df['Arrears'] == uni_arr[i]]
            temp = list(temp["Stud.Name"])
            st_name = ''
            for j in temp:
                st_name += '\n'+j
            doc.add_paragraph(st_name)
            dic[key] = temp
    doc.save('z_storage\Overall Performance.docx')
    
    context = {'image': encoded_image,"dict":dic}
    return render(request, 'overall_per.html', context)

#Subject Perforamce Function
def subject_performance(request):
    data = pd.read_csv('E:\sra\Student_Result.csv')
    sub_name = list(data.columns)
    sub_name.remove('Reg.Number')
    sub_name.remove('Stud.Name')
    sub_name.remove('Unnamed: 0')
    su_name = {'HS3151': 'Professional English - I', 'MA3151': 'Matrices and Calculus', 'PH3151': 'Engineering Physics', 'CY3151': 'Engineering Chemistry', 'GE3151': 'Problem Solving and Python Programming', 'GE3152': 'தமிழர்மரபு /Heritage of Tamils', 'GE3171': 'Problem Solving and Python Programming Laboratory', 'BS3171': 'Physics and Chemistry Laboratory', 'GE3172': 'English Laboratory $', 'HS3251': 'Professional English - II', 'MA3251': 'Statistics and Numerical Methods ', 'PH3256': 'Physics for Information Science', 'BE3251': 'Basic Electrical and Electronics  Engineering', 'GE3251': 'Engineering Graphics', 'AD3251': ' Data Structures Design', 'GE3252': 'தமிழரும் தொழில்நுட்பமும் /Tamils and Technology', 'GE3271': 'Engineering Practices Laboratory', 'AD3271': 'Data Structures Design Laboratory ', 'GE3272': 'Communication Laboratory /  Foreign Language $', 'CS3351': 'Digital Principles and  Computer Organization', 'AD3391': 'Database Design and  Management', 'AD3351': 'Design and Analysis of Algorithms', 'AD3301': 'Data Exploration and  Visualization', 'AL3391': 'Artificial Intelligence', 'AD3381': 'Database Design and  Management Laboratory', 'AD3311': 'Artificial Intelligence  Laboratory', 'GE3361': 'Professional Development$ ', 'MA3391': 'Probability and Statistics', 'AL3452': 'Operating Systems', 'AL3451': 'Machine Learning', 'AD3491': 'Fundamentals of Data  Science and Analytics', 'CS3591': 'Computer Networks', 'GE3451': 'Environmental Sciences and Sustainability', 'AD3411': 'Data Science and Analytics  Laboratory', 'AL3461': 'Machine Learning Laboratory', 'AD3501': ' Deep Learning', 'CW3551': 'Data and Information  Security', 'CS3551': 'Distributed Computing', 'CCS334': ' Big Data Analytics', 'sel': 'Professional Elective VI', 'AD3511': 'Deep Learning Laboratory', 'AD3512': 'Summer internship', 'CS3691': 'Embedded Systems and IoT', 'GE3791': 'Human Values and  Ethics', 'AD3811': 'Project Work / Internship', 'MA3354': 'Discrete Mathematics', 'CW3301': 'Fundamentals of Economics', 'CS3391': 'Object Oriented Programming', 'CW3311': 'Business Communication  Laboratory I', 'CS3381': 'Object Oriented Programming Laboratory', 'CS3492': 'Database Management Systems', 'CW3401': 'Introduction to Business  Systems', 'CS3481': 'Database Management Systems Laboratory', 'AD3461': 'Machine Learning Laboratory', 'CW3411': 'Business Communication  Laboratory II', 'CW3501': 'Fundamentals of  Management', 'CW3511': 'Summer Internship', 'CCW331': 'Business Analytics', 'CCS356': 'Object Oriented Software Engineering', 'CW3611': 'Project Work /  Internship', 'CS3251': 'Programming in C', 'CS3271': 'Programming in C Laboratory', 'CS3352': 'Foundations of Data Science', 'CD3291': 'Data Structures and Algorithms', 'CD3281': 'Data Structures and Algorithms Laboratory', 'CS3361': 'Data Science Laboratory', 'CS3452': 'Theory of Computation', 'CS3491': 'Artificial Intelligence and  Machine Learning', 'IT3401': 'Web Essentials', 'CS3451': 'Introduction to Operating Systems', 'CS3461': 'Operating Systems Laboratory', 'IT3501': 'Full Stack Web Development', 'IT3511': 'Full Stack Web Development Laboratory', 'IT3681': 'Mobile Applications Development Laboratory', 'IT3711': 'Summer internship', 'IT3811': 'Project Work /  Internship', 'CS3301': 'Data Structures ', 'CS3311': 'Data Structures Laboratory', 'CS3401': 'Algorithms', 'CS3501': 'Compiler Design', 'CB3491': 'Cryptography and Cyber Security', 'CS3711': 'Summer internship', 'CS3811': 'Project Work/Internship', 'PH3254': 'Physics for Electronics  Engineering', 'BE3254': 'Electrical and Instrumentation  Engineering', 'EC3251': 'Circuit Analysis', 'EC3271': 'Circuits Analysis Laboratory', 'MA3355': 'Random Processes and  Linear Algebra', 'CS3353': 'C Programming and Data  Structures', 'EC3354': 'Signals and Systems', 'EC3353': 'Electronic Devices and  Circuits', 'EC3351': 'Control Systems', 'EC3352': 'Digital Systems Design', 'EC3361': 'Electronic Devices and  Circuits Laboratory', 'CS3362': 'C Programming and Data  Structures Laboratory', 'EC3452': 'Electromagnetic Fields', 'EC3401': 'Networks and Security', 'EC3451': 'Linear Integrated Circuits', 'EC3492': 'Digital Signal Processing', 'EC3491': 'Communication Systems', 'EC3461': 'Communication Systems  Laboratory', 'EC3462': 'Linear Integrated Circuits  Laboratory', 'EC3501': 'Wireless Communication', 'EC3552': 'VLSI and Chip Design', 'EC3551': 'Transmission lines and RF  Systems', 'EC3561': 'VLSI Laboratory', 'ET3491': 'Embedded Systems and  IOT Design', 'EC3711': 'Summer internship', 'EC3811': 'Project Work/Internship', 'PH3202': 'Physics for Electrical  Engineering', 'BE3255': 'Basic Civil and Mechanical  Engineering', 'EE3251': 'Electric Circuit Analysis', 'EE3271': 'Electric Circuits Laboratory', 'MA3303': 'Probability and Complex  Functions', 'EE3301': 'Electromagnetic Fields', 'EE3302': 'Digital Logic Circuits', 'EC3301': 'Electron Devices and Circuits', 'EE3303': 'Electrical Machines - I', 'EC3311': 'Electronic Devices and  Circuits Laboratory', 'EE3311': 'Electrical Machines  Laboratory'}
    lst = [na+' - '+su_name[na]  for na in sub_name]
    reg = list(data['Reg.Number'])
    #print(lst)

    if request.method == 'POST':
        name = request.POST.get('subject_name')
        name = name.split('-')
        name=name[0].strip()
        #print(name)
        uni_score = list(data[name].unique())
        uni_score.sort(reverse=True)

        count = []
        for i in range (0,len(uni_score)):
            if uni_score[i] == 0:
                count.append("\n The "+str((data[name] == 0).sum())+" Students  got fail in the "+name+" Subject.")
            elif uni_score[i] == 10:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students got 'O' grade in the "+name+" Subject.")
            elif uni_score[i] == 9:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students got 'A+' grade in the "+name+" Subject.")
            elif uni_score[i] == 8:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students got 'A' grade in the "+name+" Subject.")
            elif uni_score[i] == 7:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students got 'B+' grade in the "+name+" Subject.")
            elif uni_score[i] == 6:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students got 'B' grade in the "+name+" Subject.")
            elif uni_score[i] == 5:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students got 'Just Pass' grade in the "+name+" Subject.")
            elif uni_score[i] == 1:
                count.append("\n The "+str((data[name] == uni_score[i]).sum())+" Students was absent in the "+name+" Subject.")

        plt.figure(figsize=(10,8))
        sub = sns.countplot(x=data[name])
        sub.set_title("Subject Performance")
        sub.set_ylabel('Students count')
        sub.set_xlabel('Grade')
        for p in sub.patches:
            sub.annotate('{:.0f}'.format(p.get_height()),(p.get_x() + p.get_width() / 2, p.get_height()),ha = 'center', va = 'center', xytext = (0, 10),textcoords = 'offset points')
        buf = BytesIO()
        plt.savefig(buf, format='png')
        #plt.savefig('z_storage\Subject_Performance_Bar.png')
        plt.close()
        fig = base64.b64encode(buf.getvalue()).decode('utf-8').replace('\n', '')

        label = []
        for i in range (0,len(uni_score)):
            if uni_score[i] == 0:
                label.append(str((data[name] == 0).sum())+" Students  got fail")
            elif uni_score[i] == 10:
                label.append(str((data[name] == uni_score[i]).sum())+" Students got 'O' grade")
            elif uni_score[i] == 9:
                label.append(str((data[name] == uni_score[i]).sum())+" Students got 'A+' grade")
            elif uni_score[i] == 8:
                label.append(str((data[name] == uni_score[i]).sum())+" Students got 'A' grade")
            elif uni_score[i] == 7:
                label.append(str((data[name] == uni_score[i]).sum())+" Students got 'B+' grade")
            elif uni_score[i] == 6:
                label.append(str((data[name] == uni_score[i]).sum())+" Students got 'B' grade.")
            elif uni_score[i] == 5:
                label.append(str((data[name] == uni_score[i]).sum())+" Students got 'Just Pass' grade")
            elif uni_score[i] == 1:
                label.append(str((data[name] == uni_score[i]).sum())+" Students was absent")

        score_count = []
        for i in uni_score:
            co = (data[name] == i).sum()
            score_count.append(co)
            
        plt.figure(figsize=(10,8))
        plt.pie(score_count,labels=label,autopct='%1.1f%%')
        plt.title('Subject Performance')
        
        #plt.tight_layout()
        pie = BytesIO()
        plt.savefig(pie, format='png')
        plt.close()
        pie = base64.b64encode(pie.getvalue()).decode('utf-8').replace('\n', '')

        doc = docx.Document()
        doc.add_heading(name+' - '+su_name[name]+" Subject Performance", 0)
        doc.add_heading('Performance Chart:',level=2)
        doc.add_picture(io.BytesIO(base64.b64decode(fig)),width=Inches(6))
        doc.add_picture(io.BytesIO(base64.b64decode(pie)),width=Inches(6))
        doc.add_heading('Performance Summary:',level=2)

        det=''
        for sen in count:
            det += sen
        doc.add_paragraph(det)

        dic ={}
        score_dic = {10:'O',9:'A+',8:'A',7:'B+',6:'B',5:'C',0:'get Fail',1:'Absent'}
        for i in range (0,len(uni_score)):
            key = "Student Have "+score_dic[uni_score[i]]+" Grade:"
            doc.add_heading(key,level=3)
            temp=data[data[name] == uni_score[i]]
            temp = list(temp["Stud.Name"])
            dic[key] = temp
            stu_name = ''
            for tp in temp:
                stu_name += '\n'+tp
            doc.add_paragraph(stu_name)

        doc.save('z_storage\Subject Performance.docx')
        return render(request, 'sub_per.html', {'count': count, 'img': fig,'sub_name':lst,'name':name,'pie':pie,'dict':dic})
    
    return render(request, 'sub_per.html', {'sub_name': lst})

#Student Performance Function
def student_performance(request):
    data = pd.read_csv('E:\sra\Student_Result.csv')
    sub_name = list(data.columns)
    sub_name.remove('Reg.Number')
    sub_name.remove('Stud.Name')
    sub_name.remove('Unnamed: 0')
    reg = list(data['Reg.Number'])
    if request.method == 'POST':
        d = int(request.POST.get('number_input'))

        at = data[data['Reg.Number'] == d]

        count = []
        for i in range(0, len(sub_name)):
            gra = at.at[int(at.index.values), str(sub_name[i])]
            if gra == 0:
                count.append("\nStudent is get arrear in " + str(sub_name[i]) + " subject.")
            elif gra == 10:
                count.append("\nStudent is get 'O' grade in " + str(sub_name[i]) + " subject.")
            elif gra == 9:
                count.append("\nStudent is get 'A+' grade in " + str(sub_name[i]) + " subject.")
            elif gra == 8:
                count.append("\nStudent is get 'A' grade in " + str(sub_name[i]) + " subject.")
            elif gra == 7:
                count.append("\nStudent is get 'B+' grade in " + str(sub_name[i]) + " subject.")
            elif gra == 6:
                count.append("\nStudent is get 'B' grade in " + str(sub_name[i]) + " subject.")
            elif gra == 5:
                count.append("\nStudent is get 'C' grade in " + str(sub_name[i]) + " subject.")
            elif gra == 1:
                count.append("\nStudent is get absent in " + str(sub_name[i]) + " subject.")

        dfm = at.melt(id_vars='Stud.Name')
        dfm.drop(0, axis=0, inplace=True)
        df = dfm.iloc[2:]
        fig = plt.figure(figsize=(10, 8))
        sub = fig.add_subplot(1, 1, 1)
        sub = sns.barplot(x=df['variable'],y=dfm['value'])
        for p in sub.patches:
            sub.annotate('{:.0f}'.format(p.get_height()),
                          (p.get_x() + p.get_width() / 2, p.get_height()),
                          ha='center', va='center', xytext=(0, 10), textcoords='offset points')
        st_name =  str(at.at[int(at.index.values), 'Stud.Name'])
        sub.set_title(st_name)
        sub.set_xlabel("Subjects")
        sub.set_ylabel("Grade")

        canvas = FigureCanvas(fig)
        buffer = io.BytesIO()
        canvas.print_png(buffer)
        image_data = base64.b64encode(buffer.getvalue()).decode('utf-8')
        plt.close()
        doc = docx.Document()
        doc.add_heading(st_name, 0)
        doc.add_heading('Performance Chart:', level=2)
        doc.add_picture(io.BytesIO(base64.b64decode(image_data)),width=Inches(6))
        doc.add_heading('Performance Details:', level=2)

        det=''
        for sen in count:
            det += sen
        doc.add_paragraph(det)

        
        doc.save('z_storage\Student Performance.docx')

        return render(request, 'stu_per.html', {
            'image_data': image_data,
            'summary': count,
            'reg':reg,
            'reg_no':d,
            'stu_name':st_name

        })
    return render(request, 'stu_per.html',{'reg':reg})


# GPA Calculation Function
def gpa_performance(request):
    data = pd.read_csv('E:\sra\Student_Result.csv')
    fine ={'HS3151': 3.0, 'MA3151': 4.0, 'PH3151': 3.0, 'CY3151': 3.0, 'GE3151': 3.0, 'GE3152': 1.0, 'GE3171': 2.0, 'BS3171': 2.0, 'GE3172': 1.0, 'HS3251': 2.0, 'MA3251': 4.0, 'PH3256': 3.0, 'BE3251': 3.0, 'GE3251': 4.0, 'AD3251': 3.0, 'GE3252': 1.0, 'GE3271': 2.0, 'AD3271': 2.0, 'GE3272': 2.0, 'CS3351': 4.0, 'AD3391': 3.0, 'AD3351': 4.0, 'AD3301': 4.0, 'AL3391': 3.0, 'AD3381': 1.5, 'AD3311': 1.5, 'GE3361': 1.0, 'MA3391': 4.0, 'AL3452': 4.0, 'AL3451': 3.0, 'AD3491': 3.0, 'CS3591': 4.0, 'GE3451': 2.0, 'AD3411': 2.0, 'AL3461': 2.0, 'AD3501': 3.0, 'CW3551': 3.0, 'CS3551': 3.0, 'CCS334': 3.0, 'sel': 3.0, 'AD3511': 2.0, 'AD3512': 2.0, 'CS3691': 4.0, 'GE3791': 2.0, 'AD3811': 10.0, 'MA3354': 4.0, 'CW3301': 3.0, 'CS3391': 3.0, 'CW3311': 1.5, 'CS3381': 1.5, 'CS3492': 3.0, 'CW3401': 3.0, 'CS3481': 1.5, 'AD3461': 2.0, 'CW3411': 1.5, 'CW3501': 3.0, 'CW3511': 2.0, 'CCW331': 3.0, 'CCS356': 4.0, 'CW3611': 10.0, 'CS3251': 3.0, 'CS3271': 2.0, 'CS3352': 3.0, 'CD3291': 3.0, 'CD3281': 2.0, 'CS3361': 2.0, 'CS3452': 3.0, 'CS3491': 4.0, 'IT3401': 3.0, 'CS3451': 3.0, 'CS3461': 1.5, 'IT3501': 3.0, 'IT3511': 2.0, 'IT3681': 1.5, 'IT3711': 2.0, 'IT3811': 0, 'CS3301': 3.0, 'CS3311': 1.5, 'CS3401': 4.0, 'CS3501': 4.0, 'CB3491': 3.0, 'CS3711': 2.0, 'CS3811': 10.0, 'PH3254': 3.0, 'BE3254': 3.0, 'EC3251': 4.0, 'EC3271': 1.0, 'MA3355': 4.0, 'CS3353': 3.0, 'EC3354': 4.0, 'EC3353': 3.0, 'EC3351': 3.0, 'EC3352': 4.0, 'EC3361': 1.5, 'CS3362': 1.5, 'EC3452': 3.0, 'EC3401': 4.0, 'EC3451': 3.0, 'EC3492': 4.0, 'EC3491': 3.0, 'EC3461': 1.5, 'EC3462': 1.5, 'EC3501': 4.0, 'EC3552': 3.0, 'EC3551': 3.0, 'EC3561': 2.0, 'ET3491': 4.0, 'EC3711': 2.0, 'EC3811': 10.0, 'PH3202': 3.0, 'BE3255': 3.0, 'EE3251': 4.0, 'EE3271': 2.0, 'MA3303': 4.0, 'EE3301': 4.0, 'EE3302': 3.0, 'EC3301': 3.0, 'EE3303': 3.0, 'EC3311': 1.5, 'EE3311': 1.5, 'EE3401': 3.0, 'EE3402': 3.0, 'EE3403': 3.0, 'EE3404': 3.0, 'EE3405': 3.0, 'EE3411': 1.5, 'EE3412': 1.5, 'EE3413': 1.5, 'EE3501': 3.0, 'EE3591': 3.0, 'EE3503': 3.0, 'EE3511': 1.5, 'EE3512': 2.0, 'EE3601': 3.0, 'EE3602': 3.0, 'EE3611': 1.5, 'EE3701': 3.0, 'EE3811': 10.0, 'PH3201': 3.0, 'BE3252': 3.0, 'BE3272': 1.5, 'MA3351': 4.0, 'ME3351': 3.0, 'CE3301': 3.0, 'CE3302': 3.0, 'CE3303': 4.0, 'CE3351': 3.0, 'CE3361': 1.5, 'CE3311': 1.5, 'CE3401': 4.0, 'CE3402': 3.0, 'CE3403': 3.0, 'CE3404': 3.0, 'CE3405': 3.0, 'CE3411': 1.5, 'CE3412': 3.0, 'CE3413': 1.5, 'CE3501': 3.0, 'CE3502': 3.0, 'CE3503': 3.0, 'CE3511': 2.0, 'CE3512': 1.0, 'CE3601': 3.0, 'CE3602': 3.0, 'AG3601': 3.0, 'CE3611': 2.0, 'CE3701': 3.0, 'CE3702': 2.0, 'GE3752': 3.0, 'CE3811': 10.0, 'PH3251': 3.0, 'BE3271': 2.0, 'ME3391': 3.0, 'CE3391': 4.0, 'ME3392': 3.0, 'ME3393': 3.0, 'ME3381': 2.0, 'ME3382': 2.0, 'ME3491': 3.0, 'ME3451': 4.0, 'ME3492': 3.0, 'ME3493': 3.0, 'CE3491': 3.0, 'CE3481': 2.0, 'ME3461': 2.0, 'ME3591': 4.0, 'ME3592': 3.0, 'ME3511': 1.0, 'ME3581': 2.0, 'ME3691': 4.0, 'ME3681': 2.0, 'ME3682': 2.0, 'ME3791': 3.0, 'ME3792': 3.0, 'GE3792': 3.0, 'ME3781': 2.0, 'ME3711': 1.0, 'ME3811': 10.0, 'BA4101': 3.0, 'BA4102': 3.0, 'BA4103': 3.0, 'BA4104': 3.0, 'BA4105': 3.0, 'BA4106': 3.0, 'BA4111': 2.0, 'BA4112': 2.0, 'BA4201': 3.0, 'BA4202': 3.0, 'BA4203': 3.0, 'BA4204': 3.0, 'BA4205': 3.0, 'BA4206': 3.0, 'BA4207': 3.0, 'BA4211': 2.0, 'BA4212': 2.0, 'BA4301': 3.0, 'BA4302': 3.0, 'BA4311': 2.0, 'BA4312': 2.0, 'BA4411': 12.0}
    # get the sub name for search credit points
    sub_name = list(data.columns)
    del sub_name[0:3]  # del the unwanted columns

# credit points search
    crd=[]
    for i in sub_name:
         o = fine.get(i)
         crd.append(o)

#  make data frame for calculate GPA
    a = pd.DataFrame({
    "credits":crd ,
    })
    stu_GPA = []
    stu_NAME = []
    reg = list(data['Reg.Number'])
    for i in range(0,len(reg)):
        be = data[data['Reg.Number']==reg[i]]
        bfm = be.melt(id_vars='Stud.Name')
        bfm = bfm[2:]
        bfm.reset_index(drop=True, inplace=True)
        stu_score=list(bfm["value"])
        a["std_credit"] = stu_score
        sum = a.credits * a.std_credit
        a['sum'] = sum
        gpa = a['sum'].sum()/a['credits'].sum()
        o = round(gpa,3)
        stu_GPA.append(o)
        stu_NAME.append(be['Stud.Name'].iloc[0])
    
    if len(stu_GPA) > 0:
        stu_gpa = pd.DataFrame({"Stud_Name": stu_NAME, "GPA": stu_GPA})
        final_df = stu_gpa.sort_values(by=['GPA'], ascending=False)
        final_df.reset_index(drop=True, inplace=True)
        final_df.index = np.arange(1,len(final_df)+1)
        doc = docx.Document()
        doc.add_heading('GPA Ranklist', 0)
        table = doc.add_table(rows=final_df.shape[0]+1, cols=final_df.shape[1])
        for i, column_name in enumerate(final_df.columns):
            table.cell(0, i).text = column_name
        for i in range(final_df.shape[0]):
            for j in range(final_df.shape[1]):
                table.cell(i+1, j).text = str(final_df.values[i,j])
        doc.save('z_storage/GPA_Ranklist.docx')
        return render(request,'gpa_per.html',{'dataframe':final_df})


#Upload PDF File option
def upload(request):
    if request.method == 'POST':
        semester = request.POST.get('semester')
        pdf_file = request.FILES['pdf_file']
        user_input = semester.zfill(2)
        end = str(int(user_input) + 1).zfill(2)
        at = []
        pdf_path = os.path.join('E:\\sra', pdf_file.name) # update path here
        with open(pdf_path, 'wb') as f:
            for chunk in pdf_file.chunks():
                f.write(chunk)
        with open(pdf_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfFileReader(f)
            for page_num in range(pdf_reader.getNumPages()):
                page = pdf_reader.getPage(page_num)
                text = page.extract_text()
                if 'Semester No. : ' + user_input in text:
                    for i in range(page_num + 1, pdf_reader.getNumPages()):
                        next_page = pdf_reader.getPage(i)
                        next_text = next_page.extract_text()
                        if 'Semester No. : ' + end in next_text:
                            break
                    at.append(str(page_num + 1))
                    at.append(str(i + 1))
        str_page = int(at[0])
        end_page = int(at[1])

        if str_page + 1 != end_page:
            end_page -= 1

        with open(pdf_path, 'rb') as f:
            pdf_content = f.read()
        data = pd.DataFrame()
        for i in range(str_page, end_page + 1):
            df = tabula.read_pdf(BytesIO(pdf_content), pages=i)[0]
            col = list(df.columns)
            df.rename(columns={col[0]: "Unnamed"}, inplace=True)
            data = pd.concat([data, df], ignore_index=True)
        data.columns = data.iloc[0]
        data.drop(1, inplace=True)
        data.rename(columns={data.columns[0]: 'Reg.Number', data.columns[1]: 'Stud.Name'}, inplace=True)
        data.drop(0, inplace=True)
        data.reset_index(drop=True, inplace=True)
        data.replace(np.nan, 1, inplace=True)
        replace = {'O': 10, 'UA': 1, 'A+': 9, 'A': 8, 'B+': 7, 'B': 6, 'C': 5, 'U': 0}
        for i in replace.keys():
            data.replace(i, replace[i], inplace=True)
        data['Reg.Number'] = data['Reg.Number'].apply(pd.to_numeric)
        data.to_csv('Student_Result.csv', encoding='utf-8')
        data.to_excel('Student_Result.xlsx')
        sub_name = list(data.columns)
        sub_name.remove('Reg.Number')
        sub_name.remove('Stud.Name')
        os.remove(pdf_path)
        context = {'dataframe': data}
        return render(request, 'pdf_to_dataframe.html', context)

    return render(request, 'pdf_form.html')

@login_required()   
def profile(request):
    return render(request,'profile.html')

#Home Page Function
def index(request):
    return render(request,'house.html')

def register(request):
    if request.method=='POST':
        name=request.POST['username']
        email=request.POST['email']
        password1=request.POST['password1']
        password2=request.POST['password2']

        if password1 == password2:
            user = User.objects.create_user(username=name,email=email,password=password1)
            user.is_staff = True
            #user.is_superuser = True
            user.save()
            messages.success(request,'Your account has been created. You have login.')
            return redirect('login')
        else:
            messages.warning(request,'Password Mismatching...!!!')
            return redirect('register')

    else:    
        form=CreateUserForm()
        return render(request,'register.html',{'form':form})

def download_file(request):
    file_path = 'E:\sra\Student_Result.xlsx' 
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/force-download')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(file_path)
            return response
    raise Http404 

def download_stu_file(request):
    file_path = 'E:\sra\z_storage\Student Performance.docx' 
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/force-download')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(file_path)
            return response
    raise Http404 

def download_sub_file(request):
    file_path = 'E:\sra\z_storage\Subject Performance.docx' 
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/force-download')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(file_path)
            return response
    raise Http404 

def download_overall_file(request):
    file_path = 'E:\sra\z_storage\Overall Performance.docx' 
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/force-download')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(file_path)
            return response
    raise Http404 

def download_gpa_file(request):
    file_path = 'E:\sra\z_storage\GPA_Ranklist.docx' 
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/force-download')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(file_path)
            return response
    raise Http404 